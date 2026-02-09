"""
Question Paper Generator

Generates formatted exam papers (PDF/Word) from question bank files.
Supports embedding geometry figures and matching ICSE exam paper format.

Usage:
    # Generate PDF paper
    python paper_generator.py --input Geometry_Questions.txt --output exam_paper.pdf --format pdf
    
    # Generate with specific criteria
    python paper_generator.py --input Geometry_Questions.txt \\
        --topics "Circles,Similarity" \\
        --total-marks 80 \\
        --output geometry_paper.pdf
    
    # Generate Word document
    python paper_generator.py --input Geometry_Questions.txt --output exam_paper.docx --format docx
"""

import argparse
import os
import re
import io
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Tuple, Any

# Optional imports with fallbacks
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Image as RLImage, PageBreak, KeepTogether
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = Any

try:
    from geometry_schema import FigureParser, GeometryFigure
    from figure_renderer import FigureRenderer, RenderConfig
    GEOMETRY_AVAILABLE = True
except ImportError:
    GEOMETRY_AVAILABLE = False


# ============================================================================
# Data Models
# ============================================================================

@dataclass
class Question:
    """Represents a parsed question from the question bank."""
    
    number: str                      # Original question number (e.g., "Q5", "Q1(ii)")
    marks: int                       # Marks for the question
    difficulty: str                  # easy, medium, hard
    text: str                        # Full question text
    topic: str                       # Topic name
    subtopic: Optional[str] = None   # Subtopic classification
    source: Optional[str] = None     # Source paper reference
    is_mcq: bool = False             # Is this an MCQ?
    figure_block: Optional[str] = None  # Raw [FIGURE]...[/FIGURE] content
    sub_parts: List[Dict] = field(default_factory=list)  # Sub-parts for multi-part questions


@dataclass
class Section:
    """Represents a section in the exam paper."""
    
    name: str                        # e.g., "Section A", "Section B"
    description: str                 # e.g., "Attempt all questions (10 × 1 = 10 marks)"
    questions: List[Question] = field(default_factory=list)
    total_marks: int = 0


@dataclass
class ExamPaper:
    """Represents a complete exam paper."""
    
    title: str = "Mathematics Examination"
    subtitle: str = "ICSE Class 10"
    duration: str = "2 hours"
    total_marks: int = 80
    date: str = ""
    instructions: List[str] = field(default_factory=list)
    sections: List[Section] = field(default_factory=list)
    
    def __post_init__(self):
        if not self.date:
            self.date = datetime.now().strftime("%B %Y")
        if not self.instructions:
            self.instructions = [
                "Answers to this paper must be written on the paper provided separately.",
                "You will not be allowed to write during the first 15 minutes.",
                "This time is to be spent reading the question paper.",
                "The time given at the head of this paper is the time allowed for writing the answers.",
                "Attempt all questions from Section A and any four questions from Section B.",
                "All working, including rough work, must be clearly shown.",
                "Mathematical tables are provided."
            ]


# ============================================================================
# Question Bank Parser
# ============================================================================

class QuestionBankParser:
    """Parses question bank .txt files into Question objects."""
    
    # Pattern to match question headers like "Q5 [4 marks] (medium)" or "Q1(ii) [1 mark] (easy) - MCQ"
    QUESTION_PATTERN = re.compile(
        r'^(Q\d+[a-z]?(?:\([ivxlcdm]+\)|[a-z])?)\s*\[(\d+)\s*marks?\]\s*\((\w+)\)(?:\s*-\s*MCQ)?',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern to match topic headers
    TOPIC_PATTERN = re.compile(
        r'^-+\s*\nTopic:\s*(.+?)\s*\n',
        re.MULTILINE
    )
    
    # Pattern to extract [FIGURE]...[/FIGURE] blocks
    FIGURE_PATTERN = re.compile(
        r'\[FIGURE\](.*?)\[/FIGURE\]',
        re.DOTALL | re.IGNORECASE
    )
    
    # Pattern to match source
    SOURCE_PATTERN = re.compile(r'\[Source:\s*(.+?)\]', re.IGNORECASE)
    
    # Pattern to match subtopic
    SUBTOPIC_PATTERN = re.compile(r'Subtopic:\s*(.+?)(?:\n|$)', re.IGNORECASE)
    
    # Pattern to match separator lines
    SEPARATOR_PATTERN = re.compile(r'-{10,}')

    def __init__(self):
        self.questions: List[Question] = []
        self.topics: Dict[str, List[Question]] = {}
    
    def parse_file(self, filepath: str) -> List[Question]:
        """Parse a question bank file and return list of questions."""
        
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return self.parse_content(content)
    
    def parse_content(self, content: str) -> List[Question]:
        """Parse question bank content."""
        
        self.questions = []
        self.topics = {}
        
        # Split by topic sections
        # Split by topic sections
        # Use capturing group to keep the delimiter (topic name)
        # Pattern matches: (optional separator) + "Topic:" + space + (topic name) + newline
        # We use a lookahead or just split on the whole block
        # Improved pattern: matches optional dashes, then Topic: <name>
        parts = re.split(r'(?:^|\n)(?:-{10,}\s*\n)?Topic:\s*(.+?)\s*\n', content)
        
        # The first part is content before the first topic separator (usually header or "General")
        if parts[0].strip():
            self._parse_questions_in_section(parts[0], "General")
            
        # Subsequent parts are (topic_name, section_content) pairs
        for i in range(1, len(parts), 2):
            if i + 1 < len(parts):
                topic_name = parts[i].strip()
                section_content = parts[i+1]
                self._parse_questions_in_section(section_content, topic_name)
        
        return self.questions
    
    def _parse_questions_in_section(self, section: str, topic: str):
        """Parse questions within a topic section."""
        
        # Find all question matches
        matches = list(self.QUESTION_PATTERN.finditer(section))
        
        for i, match in enumerate(matches):
            # Get the content until the next question or separator
            start = match.end()
            if i + 1 < len(matches):
                end = matches[i + 1].start()
            else:
                # Find the separator
                sep_match = re.search(r'-{20,}', section[start:])
                end = start + sep_match.start() if sep_match else len(section)
            
            question_content = section[start:end].strip()
            
            # Parse question details
            q_number = match.group(1)
            q_marks = int(match.group(2))
            q_difficulty = match.group(3).lower()
            is_mcq = 'MCQ' in match.group(0).upper()
            
            # Extract figure block if present
            figure_match = self.FIGURE_PATTERN.search(question_content)
            figure_block = figure_match.group(1).strip() if figure_match else None
            
            # Remove figure block from question text
            if figure_block:
                question_content = self.FIGURE_PATTERN.sub('', question_content).strip()
            
            # Extract source
            source_match = self.SOURCE_PATTERN.search(question_content)
            source = source_match.group(1).strip() if source_match else None
            if source:
                question_content = self.SOURCE_PATTERN.sub('', question_content).strip()
            
            # Extract subtopic
            subtopic_match = self.SUBTOPIC_PATTERN.search(question_content)
            subtopic = subtopic_match.group(1).strip() if subtopic_match else None
            if subtopic:
                question_content = self.SUBTOPIC_PATTERN.sub('', question_content).strip()
            
            # Clean up the question text
            question_text = question_content.strip()
            question_text = self.SEPARATOR_PATTERN.sub('', question_text).strip()
            
            question = Question(
                number=q_number,
                marks=q_marks,
                difficulty=q_difficulty,
                text=question_text,
                topic=topic,
                subtopic=subtopic,
                source=source,
                is_mcq=is_mcq,
                figure_block=figure_block
            )
            
            self.questions.append(question)
            
            # Add to topic dict
            if topic not in self.topics:
                self.topics[topic] = []
            self.topics[topic].append(question)
    
    def get_questions_by_topic(self, topic: str) -> List[Question]:
        """Get questions for a specific topic."""
        return self.topics.get(topic, [])
    
    def get_questions_by_marks(self, marks: int) -> List[Question]:
        """Get questions with specific marks."""
        return [q for q in self.questions if q.marks == marks]
    
    def get_questions_by_difficulty(self, difficulty: str) -> List[Question]:
        """Get questions by difficulty level."""
        return [q for q in self.questions if q.difficulty == difficulty.lower()]
    
    def filter_questions(
        self,
        topics: Optional[List[str]] = None,
        marks: Optional[int] = None,
        difficulty: Optional[str] = None,
        mcq_only: bool = False,
        exclude_mcq: bool = False
    ) -> List[Question]:
        """Filter questions by multiple criteria."""
        
        result = self.questions.copy()
        
        if topics:
            topics_lower = [t.lower() for t in topics]
            result = [q for q in result if q.topic.lower() in topics_lower]
        
        if marks is not None:
            result = [q for q in result if q.marks == marks]
        
        if difficulty:
            result = [q for q in result if q.difficulty == difficulty.lower()]
        
        if mcq_only:
            result = [q for q in result if q.is_mcq]
        
        if exclude_mcq:
            result = [q for q in result if not q.is_mcq]
        
        return result


# ============================================================================
# PDF Generator (using ReportLab)
# ============================================================================

class PDFPaperGenerator:
    """Generates PDF exam papers using ReportLab."""
    
    def __init__(self, render_figures: bool = True):
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab is required for PDF generation. Install with: pip install reportlab")
        
        self.render_figures = render_figures
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
        self.figure_parser = FigureParser() if GEOMETRY_AVAILABLE else None
        self.temp_images: List[str] = []
    
    def _setup_custom_styles(self):
        """Set up custom paragraph styles."""
        
        self.styles.add(ParagraphStyle(
            name='PaperTitle',
            parent=self.styles['Heading1'],
            fontName='Times-Bold',
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=4
        ))
        
        self.styles.add(ParagraphStyle(
            name='PaperSubtitle',
            parent=self.styles['Normal'],
            fontName='Times-Roman',
            fontSize=12,
            alignment=TA_CENTER,
            spaceAfter=8
        ))
        
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading2'],
            fontName='Times-Bold',
            fontSize=13,
            spaceBefore=10,
            spaceAfter=4,
            textColor=colors.black,
            borderPadding=2,
            borderWidth=0.5,
            borderColor=colors.grey
        ))
        
        self.styles.add(ParagraphStyle(
            name='QuestionNumber',
            parent=self.styles['Normal'],
            fontSize=11,
            fontName='Times-Bold',
            spaceBefore=6,
            spaceAfter=2
        ))
        
        self.styles.add(ParagraphStyle(
            name='QuestionText',
            parent=self.styles['Normal'],
            fontName='Times-Roman',
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=4,
            leftIndent=0
        ))
        
        self.styles.add(ParagraphStyle(
            name='Instructions',
            parent=self.styles['Normal'],
            fontName='Times-Roman',
            fontSize=10,
            leftIndent=15,
            spaceAfter=2
        ))
    
    def generate(self, paper: ExamPaper, output_path: str):
        """Generate a PDF exam paper."""
        
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        story = []
        
        # Header
        story.extend(self._build_header(paper))
        
        # Instructions
        story.extend(self._build_instructions(paper))
        
        # Sections
        for section in paper.sections:
            story.extend(self._build_section(section))
        
        # Build document
        doc.build(story)
        
        # Cleanup temp images
        self._cleanup_temp_images()
        
        print(f"Generated PDF: {output_path}")
    
    def _build_header(self, paper: ExamPaper) -> List:
        """Build the paper header."""
        
        elements = []
        
        elements.append(Paragraph(paper.title, self.styles['PaperTitle']))
        elements.append(Paragraph(paper.subtitle, self.styles['PaperSubtitle']))
        
        # Duration and marks table
        info_data = [
            [f"Time: {paper.duration}", f"Maximum Marks: {paper.total_marks}"]
        ]
        info_table = Table(info_data, colWidths=[8*cm, 8*cm])
        info_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
        ]))
        elements.append(info_table)
        elements.append(Spacer(1, 12))
        
        return elements
    
    def _build_instructions(self, paper: ExamPaper) -> List:
        """Build the instructions section."""
        
        elements = []
        
        elements.append(Paragraph("<b>General Instructions:</b>", self.styles['Normal']))
        elements.append(Spacer(1, 6))
        
        for i, instruction in enumerate(paper.instructions, 1):
            elements.append(Paragraph(
                f"{i}. {instruction}",
                self.styles['Instructions']
            ))
        
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("_" * 80, self.styles['Normal']))
        elements.append(Spacer(1, 12))
        
        return elements
    
    def _build_section(self, section: Section) -> List:
        """Build a section of the paper."""
        
        elements = []
        
        # Section header
        elements.append(Paragraph(
            f"<b>{section.name}</b>",
            self.styles['SectionHeader']
        ))
        elements.append(Paragraph(section.description, self.styles['Normal']))
        elements.append(Spacer(1, 12))
        
        # Questions
        for i, question in enumerate(section.questions, 1):
            elements.extend(self._build_question(question, i))
        
        return elements
    
    def _build_question(self, question: Question, display_number: int) -> List:
        """Build a single question."""
        
        elements = []
        
        # Question number and marks
        marks_text = f"[{question.marks}]" if question.marks else ""
        elements.append(Paragraph(
            f"<b>Question {display_number}</b> {marks_text}",
            self.styles['QuestionNumber']
        ))
        
        # Question text
        # Escape HTML special characters
        text = question.text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        text = text.replace('\n', '<br/>')
        elements.append(Paragraph(text, self.styles['QuestionText']))
        
        # Render figure if present and rendering is enabled
        if question.figure_block:
            if self.render_figures and GEOMETRY_AVAILABLE:
                figure_elements = self._render_figure(question.figure_block)
                elements.extend(figure_elements)
            else:
                # Add placeholder text for figure
                elements.append(Spacer(1, 6))
                elements.append(Paragraph(
                    "<i>[Diagram: See figure description in question]</i>",
                    self.styles['QuestionText']
                ))
        
        elements.append(Spacer(1, 6))
        
        return elements
    
    def _render_figure(self, figure_block: str) -> List:
        """Render a geometry figure and return ReportLab elements."""
        
        elements = []
        
        try:
            figure = self.figure_parser.parse(figure_block)
            
            # Create a temporary image
            renderer = FigureRenderer(RenderConfig(figsize=(5, 5), dpi=150))
            renderer.render(figure)
            
            # Render to memory
            img_buffer = io.BytesIO()
            renderer.save_png(img_buffer)
            renderer.close()
            img_buffer.seek(0)
            
            # Add image to document
            elements.append(Spacer(1, 6))
            elements.append(RLImage(img_buffer, width=3*inch, height=3*inch))
            elements.append(Spacer(1, 6))
            
        except Exception as e:
            print(f"Warning: Could not render figure: {e}")
        
        return elements
    
    def _cleanup_temp_images(self):
        """Clean up temporary image files."""
        for path in self.temp_images:
            try:
                os.unlink(path)
            except:
                pass
        self.temp_images = []


# ============================================================================
# Word Generator (using python-docx)
# ============================================================================

class WordPaperGenerator:
    """Generates Word document exam papers using python-docx."""
    
    def __init__(self, render_figures: bool = True):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word generation. Install with: pip install python-docx")
        
        self.render_figures = render_figures
        self.figure_parser = FigureParser() if GEOMETRY_AVAILABLE else None
        self.temp_images: List[str] = []
    
    def generate(self, paper: ExamPaper, output_path: str):
        """Generate a Word document exam paper."""
        
        doc = Document()
        
        # Set up styles
        self._setup_styles(doc)
        
        # Header
        self._build_header(doc, paper)
        
        # Instructions
        self._build_instructions(doc, paper)
        
        # Sections
        for section in paper.sections:
            self._build_section(doc, section)
        
        # Save document
        doc.save(output_path)
        
        # Cleanup temp images
        self._cleanup_temp_images()
        
        print(f"Generated Word document: {output_path}")
    
    def _setup_styles(self, doc: Document):
        """Set up document styles."""
        
        # Modify Normal style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
    
    def _build_header(self, doc: Document, paper: ExamPaper):
        """Build the paper header."""
        
        # Title
        title = doc.add_heading(paper.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph(paper.subtitle)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Duration and marks
        info = doc.add_paragraph()
        info.add_run(f"Time: {paper.duration}").bold = True
        info.add_run(" " * 40)
        info.add_run(f"Maximum Marks: {paper.total_marks}").bold = True
        
        doc.add_paragraph()
    
    def _build_instructions(self, doc: Document, paper: ExamPaper):
        """Build the instructions section."""
        
        doc.add_paragraph("General Instructions:").bold = True
        
        for i, instruction in enumerate(paper.instructions, 1):
            doc.add_paragraph(f"{i}. {instruction}", style='List Number')
        
        doc.add_paragraph("_" * 70)
        doc.add_paragraph()
    
    def _build_section(self, doc: Document, section: Section):
        """Build a section of the paper."""
        
        # Section header
        heading = doc.add_heading(section.name, level=2)
        doc.add_paragraph(section.description)
        
        # Questions
        for i, question in enumerate(section.questions, 1):
            self._build_question(doc, question, i)
    
    def _build_question(self, doc: Document, question: Question, display_number: int):
        """Build a single question."""
        
        # Question number and marks
        p = doc.add_paragraph()
        marks_text = f" [{question.marks}]" if question.marks else ""
        p.add_run(f"Question {display_number}{marks_text}").bold = True
        
        # Question text
        doc.add_paragraph(question.text)
        
        # Render figure if present and rendering is enabled
        if question.figure_block:
            if self.render_figures and GEOMETRY_AVAILABLE:
                self._render_figure(doc, question.figure_block)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run("[Diagram: See figure description in question]").italic = True
        
        doc.add_paragraph()
    
    def _render_figure(self, doc: Document, figure_block: str):
        """Render a geometry figure and add to document."""
        
        try:
            figure = self.figure_parser.parse(figure_block)
            
            # Create a temporary image
            renderer = FigureRenderer(RenderConfig(figsize=(5, 5), dpi=150))
            renderer.render(figure)
            
            # Render to memory
            img_buffer = io.BytesIO()
            renderer.save_png(img_buffer)
            renderer.close()
            img_buffer.seek(0)
            
            # Add image to document
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_buffer, width=Inches(3))
            
        except Exception as e:
            print(f"Warning: Could not render figure: {e}")
    
    def _cleanup_temp_images(self):
        """Clean up temporary image files."""
        for path in self.temp_images:
            try:
                os.unlink(path)
            except:
                pass
        self.temp_images = []


# ============================================================================
# Paper Builder
# ============================================================================

class PaperBuilder:
    """
    High-level interface for building exam papers from question banks.
    
    Handles question selection, section distribution, and paper generation.
    """
    
    # Default ICSE paper structure
    DEFAULT_STRUCTURE = {
        "Section A": {
            "description": "Attempt all questions from this Section. (10 × 1 = 10 marks)",
            "question_count": 10,
            "marks_per_question": 1,
            "mcq_only": True
        },
        "Section B": {
            "description": "Attempt any four questions from this Section. (5 × 6 = 30 marks + 4 × 10 = 40 marks)",
            "subsections": [
                {"count": 5, "marks": 6},
                {"count": 4, "marks": 10}
            ]
        }
    }
    
    def __init__(self):
        self.parser = QuestionBankParser()
        self.questions: List[Question] = []
    
    def load_questions(self, *filepaths: str) -> 'PaperBuilder':
        """Load questions from one or more question bank files."""
        
        for filepath in filepaths:
            if os.path.exists(filepath):
                questions = self.parser.parse_file(filepath)
                self.questions.extend(questions)
                print(f"Loaded {len(questions)} questions from {filepath}")
            else:
                print(f"Warning: File not found: {filepath}")
        
        return self
    
    def build_paper(
        self,
        title: str = "Mathematics Examination",
        subtitle: str = "ICSE Class 10",
        total_marks: int = 80,
        topics: Optional[List[str]] = None,
        auto_select: bool = True
    ) -> ExamPaper:
        """
        Build an exam paper from loaded questions.
        
        Args:
            title: Paper title
            subtitle: Paper subtitle
            total_marks: Total marks for the paper
            topics: List of topics to include (None = all topics)
            auto_select: Automatically select questions based on marks distribution
            
        Returns:
            ExamPaper object ready for generation
        """
        
        paper = ExamPaper(
            title=title,
            subtitle=subtitle,
            total_marks=total_marks
        )
        
        # Filter by topics if specified
        available = self.questions
        if topics:
            topics_lower = [t.lower() for t in topics]
            available = [q for q in available if any(t in q.topic.lower() for t in topics_lower)]
        
        if auto_select:
            paper.sections = self._auto_select_sections(available, total_marks)
        
        return paper
    
    def _auto_select_sections(self, available: List[Question], total_marks: int) -> List[Section]:
        """Automatically select questions and organize into sections."""
        
        sections = []
        
        # Section A: MCQs (1 mark each)
        mcqs = [q for q in available if q.is_mcq or q.marks == 1]
        section_a = Section(
            name="Section A",
            description="Attempt all questions. (Multiple Choice Questions)",
            total_marks=10
        )
        section_a.questions = mcqs[:10] if len(mcqs) >= 10 else mcqs
        sections.append(section_a)
        
        # Section B: Short/Long answer
        non_mcqs = [q for q in available if not q.is_mcq and q.marks > 1]
        
        # Split by marks
        short_answer = [q for q in non_mcqs if 3 <= q.marks <= 6]
        long_answer = [q for q in non_mcqs if q.marks >= 8]
        
        section_b = Section(
            name="Section B",
            description="Attempt any four questions from this Section.",
            total_marks=total_marks - 10
        )
        
        # Add short answer questions
        section_b.questions.extend(short_answer[:5])
        
        # Add long answer questions
        section_b.questions.extend(long_answer[:4])
        
        sections.append(section_b)
        
        return sections
    
    def generate_pdf(self, paper: ExamPaper, output_path: str, render_figures: bool = True):
        """Generate a PDF from the exam paper."""
        
        generator = PDFPaperGenerator(render_figures=render_figures)
        generator.generate(paper, output_path)
    
    def generate_docx(self, paper: ExamPaper, output_path: str, render_figures: bool = True):
        """Generate a Word document from the exam paper."""
        
        generator = WordPaperGenerator(render_figures=render_figures)
        generator.generate(paper, output_path)
    
    def generate(self, paper: ExamPaper, output_path: str, format: str = "pdf", render_figures: bool = True):
        """Generate paper in specified format."""
        
        if format.lower() == "pdf":
            self.generate_pdf(paper, output_path, render_figures=render_figures)
        elif format.lower() in ("docx", "word"):
            self.generate_docx(paper, output_path, render_figures=render_figures)
        else:
            raise ValueError(f"Unsupported format: {format}")


# ============================================================================
# CLI Interface
# ============================================================================

def main():
    """Command-line interface for paper generation."""
    
    parser = argparse.ArgumentParser(
        description="Generate formatted exam papers from question banks.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python paper_generator.py --input Geometry_Questions.txt --output exam.pdf
  python paper_generator.py --input questions.txt --topics "Circles,Similarity" --output geo.pdf
  python paper_generator.py --input questions.txt --output exam.docx --format docx
        """
    )
    
    parser.add_argument(
        '--input', '-i',
        nargs='+',
        help='Input question bank file(s)'
    )
    
    parser.add_argument(
        '--output', '-o',
        help='Output file path (PDF or DOCX)'
    )
    
    parser.add_argument(
        '--format', '-f',
        choices=['pdf', 'docx', 'word'],
        default='pdf',
        help='Output format (default: pdf)'
    )
    
    parser.add_argument(
        '--title',
        default='Mathematics Examination',
        help='Paper title'
    )
    
    parser.add_argument(
        '--subtitle',
        default='ICSE Class 10',
        help='Paper subtitle'
    )
    
    parser.add_argument(
        '--total-marks',
        type=int,
        default=80,
        help='Total marks for the paper (default: 80)'
    )
    
    parser.add_argument(
        '--topics',
        help='Comma-separated list of topics to include'
    )
    
    parser.add_argument(
        '--check-deps',
        action='store_true',
        help='Check available dependencies and exit'
    )
    
    parser.add_argument(
        '--no-figures',
        action='store_true',
        help='Skip figure rendering (questions with [FIGURE] blocks will show placeholder text)'
    )
    
    args = parser.parse_args()
    
    # Check dependencies
    if args.check_deps:
        print("Dependency Check:")
        print(f"  ReportLab (PDF): {'✓ Available' if REPORTLAB_AVAILABLE else '✗ Not installed'}")
        print(f"  python-docx (Word): {'✓ Available' if DOCX_AVAILABLE else '✗ Not installed'}")
        print(f"  Geometry Schema: {'✓ Available' if GEOMETRY_AVAILABLE else '✗ Not available'}")
        return
    
    # Validate required arguments
    if not args.input:
        parser.error("--input is required when generating papers")
    if not args.output:
        parser.error("--output is required when generating papers")
    
    # Parse topics
    topics = None
    if args.topics:
        topics = [t.strip() for t in args.topics.split(',')]
    
    # Build and generate paper
    builder = PaperBuilder()
    builder.load_questions(*args.input)
    
    paper = builder.build_paper(
        title=args.title,
        subtitle=args.subtitle,
        total_marks=args.total_marks,
        topics=topics
    )
    
    # Auto-detect format from extension if not specified
    format = args.format
    if args.output.endswith('.docx'):
        format = 'docx'
    elif args.output.endswith('.pdf'):
        format = 'pdf'
    
    builder.generate(paper, args.output, format, render_figures=not args.no_figures)


if __name__ == "__main__":
    main()
